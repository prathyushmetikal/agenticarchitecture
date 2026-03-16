import os
import docx
import pdfplumber
import subprocess
#from langchain.tools import tool

#@tool
def extract_cv_text(file_path)->str:
    # """
    # Extracts the text content from the resume file in pdf, docx or doc format . The resume should be inside the Docs folder .

    # Args:
    # file_path : the local file path to the resume document e.g. "./Docs/just-text.pdf"

    # Returns:
    # str: the extracted plain text from the resume , or an error message if the format is unsupported or cannot be read.
    
    # """
    print(">>> Entering extract_cv_text")
    ext=os.path.splitext(file_path)[-1].lower()

    if ".docx" in ext:
        try:
            doc=docx.Document(file_path)
            text=[para.text for para in doc.paragraphs]
            return '\n'.join(text)
        except Exception as e:
            return f"Error reading the .doc file: {e}"
    elif ".pdf" in ext:
        try:
            text=[]
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text=page.extract_text()
                    if page_text:
                        text.append(page_text)
            return '\n'.join(text)
        except Exception as e:
            return f"Error reading the .pdf file: {e}"
    elif ".doc" in ext:
        try:
            temp_docx=file_path+".temp.docx"
            subprocess.run(['soffice','--headless','docx','--outdir',os.path.dirname(file_path),file_path],check=True)
            doc=docx.Document(os.path.splitext(file_path)[0]+".docx")
            text=[para.text for para in doc.paragraphs]
            os.remove(os.path.splitext(file_path)[0]+".docx")
            return '\n'.join(text)
        
        except Exception as e:
            return f"Error reading .doc file : {e}"
    else:
        return " Unsupported file format!! Please use .pdf,.doc or .docx"
    
    print(">>> Exiting extract_cv_text")
